# services/course/manager.py
import asyncio
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, and_, func, case
from typing import List, Optional, Dict, Any, Tuple
import logging
from datetime import datetime, timedelta, timezone

from ...models.course import Course, Lesson, Activity, LessonTemplate
from ...schemas.course import CourseCreate, CourseUpdate, LessonCreate
from ...core.exceptions import NotFoundException, ValidationError
from ...services.content import ContentGenerator as AIService
from ...services.optimization.query_optimizer import QueryOptimizer
from ...services.optimization.batch_processor import BatchProcessor
from ...core.memory import memory_optimized
from ...core.cache import CacheService

logger = logging.getLogger(__name__)


class CourseManager:
    def __init__(self, session: AsyncSession):
        self.session = session
        self.ai_service = AIService(session)
        self.batch_processor = BatchProcessor(session)
        self.query_optimizer = QueryOptimizer(session)
        self.cache_service = CacheService()

    @memory_optimized()
    async def track_course_usage(self, user_id: int, course_id: int, interaction_type: str = "view", section: Optional[str] = None, metadata: Optional[Dict[str, Any]] = None) -> bool:
        """Отмечает использование курса с оптимизацией запросов и отслеживанием в аналитике"""
        try:
            query = await self.query_optimizer.optimize_query(
                select(Course).where(Course.id == course_id)
            )
            result = await self.session.execute(query)
            course = result.scalar_one_or_none()

            if course:
                course.is_used = True
                course.last_used = datetime.now(timezone.utc)
                course.usage_count += 1
                await self.session.commit()

                # Отслеживаем взаимодействие с курсом в аналитике
                try:
                    from ..analytics import OptimizedAnalyticsService
                    analytics_service = OptimizedAnalyticsService(self.session)
                    await analytics_service.track_course_interaction(
                        user_id=user_id,
                        course_id=course_id,
                        interaction_type=interaction_type,
                        course_name=course.name,
                        section=section,
                        metadata=metadata or {
                            "course_type": course.type,
                            "course_level": course.level,
                            "course_subject": course.subject,
                            "is_template": course.is_template
                        }
                    )
                except Exception as e:
                    logger.error(f"Error tracking course interaction: {str(e)}")

                return True
            return False
        except Exception as e:
            logger.error(f"Error tracking course usage: {str(e)}")
            await self.session.rollback()
            return False

    @memory_optimized()
    async def update_course(self, course_id: int, course_data: CourseUpdate) -> Optional[Course]:
        """Обновляет существующий курс"""
        try:
            # Получаем существующий курс
            query = await self.query_optimizer.optimize_query(
                select(Course).where(Course.id == course_id)
            )
            result = await self.session.execute(query)
            course = result.scalar_one_or_none()

            if not course:
                logger.warning(f"Course with ID {course_id} not found for update")
                return None

            # Обновляем поля курса
            try:
                # Пробуем использовать model_dump (Pydantic v2)
                update_data = course_data.model_dump(exclude_unset=True)
            except AttributeError:
                try:
                    # Пробуем использовать dict (Pydantic v1)
                    update_data = course_data.dict(exclude_unset=True)
                except AttributeError:
                    # Если ни один метод не работает, преобразуем объект в словарь вручную
                    update_data = {k: v for k, v in course_data.__dict__.items()
                                  if not k.startswith('_') and v is not None}

            # Обновляем поля курса
            for key, value in update_data.items():
                setattr(course, key, value)

            # Обновляем дату изменения
            course.updated_at = datetime.now(timezone.utc)

            # Сохраняем изменения
            await self.session.commit()
            await self.session.refresh(course)  # Обновляем объект из базы данных

            # Логируем успешное обновление
            logger.info(f"Course with ID {course_id} successfully updated")

            # Возвращаем обновленный курс
            return course
        except Exception as e:
            logger.error(f"Error updating course: {str(e)}")
            await self.session.rollback()
            raise

    @memory_optimized()
    async def create_course(self, course_data: CourseCreate, user_id: int) -> Course:
        """Создает новый курс с оптимизированной обработкой уроков и активностей"""
        try:
            self._validate_course_data(course_data)

            # Calculate total_duration from lessons before creating the course
            total_duration = 0
            for lesson_data in course_data.lessons:
                lesson_duration = lesson_data.duration
                # Add durations of activities if they exist
                if lesson_data.activities:
                    activity_duration = sum(activity.duration for activity in lesson_data.activities)
                    if activity_duration > 0:
                        lesson_duration = activity_duration
                total_duration += lesson_duration

            # Create course with total_duration set
            try:
                # Пробуем использовать model_dump (Pydantic v2)
                course_dict = course_data.model_dump(exclude={'lessons'})
            except AttributeError:
                try:
                    # Пробуем использовать dict (Pydantic v1)
                    course_dict = course_data.dict(exclude={'lessons'})
                except AttributeError:
                    # Если ни один метод не работает, преобразуем объект в словарь вручную
                    course_dict = {k: v for k, v in course_data.__dict__.items()
                                  if not k.startswith('_') and k != 'lessons'}

            course = Course(
                **course_dict,
                creator_id=user_id,
                is_used=False,
                usage_count=0,
                total_duration=total_duration  # Set initial total_duration
            )
            self.session.add(course)
            await self.session.flush()

            # Используем BatchProcessor для создания уроков
            lessons_data = []
            recalculated_total_duration = 0

            for lesson_data in course_data.lessons:
                lesson = await self._create_lesson_with_batch(lesson_data, course.id)
                lessons_data.append(lesson)
                recalculated_total_duration += lesson.duration

            await self.batch_processor.bulk_insert(lessons_data)

            # Update total_duration with the actual calculated value
            course.total_duration = recalculated_total_duration
            await self.session.commit()

            # Reload the course with all relationships to avoid DetachedInstanceError
            from sqlalchemy import select
            from sqlalchemy.orm import selectinload

            # Create a query with explicit loading of all relationships
            query = (
                select(Course)
                .where(Course.id == course.id)
                .options(
                    selectinload(Course.lessons).selectinload(Lesson.activities)
                )
            )

            # Execute the query
            result = await self.session.execute(query)
            refreshed_course = result.scalar_one()

            # Кэшируем созданный курс
            cache_key = f"course:{refreshed_course.id}"
            await self.cache_service.cache_data(cache_key, refreshed_course, ttl=3600)

            return refreshed_course

        except Exception as e:
            logger.error(f"Error creating course: {str(e)}")
            await self.session.rollback()
            raise

    async def _create_lesson_with_batch(self, lesson_data: LessonCreate, course_id: int) -> Lesson:
        """Создает урок и активности с использованием батч-процессинга"""
        # Calculate lesson duration from activities if they exist
        lesson_duration = lesson_data.duration
        if lesson_data.activities:
            activities_duration = sum(activity.duration for activity in lesson_data.activities)
            if activities_duration > 0:
                lesson_duration = activities_duration

        # Create lesson with proper duration - exclude both activities and duration from dict
        try:
            # Пробуем использовать model_dump (Pydantic v2)
            lesson_dict = lesson_data.model_dump(exclude={'activities', 'duration'})
        except AttributeError:
            try:
                # Пробуем использовать dict (Pydantic v1)
                lesson_dict = lesson_data.dict(exclude={'activities', 'duration'})
            except AttributeError:
                # Если ни один метод не работает, преобразуем объект в словарь вручную
                lesson_dict = {k: v for k, v in lesson_data.__dict__.items()
                              if not k.startswith('_') and k not in ['activities', 'duration']}

        lesson = Lesson(
            **lesson_dict,
            course_id=course_id,
            duration=lesson_duration  # Set duration explicitly
        )
        self.session.add(lesson)
        await self.session.flush()

        # Подготавливаем активности для батч-вставки
        activities = []

        for activity_data in lesson_data.activities:
            # Create activity with proper parameters
            try:
                # Пробуем использовать model_dump (Pydantic v2)
                activity_dict = activity_data.model_dump()
            except AttributeError:
                try:
                    # Пробуем использовать dict (Pydantic v1)
                    activity_dict = activity_data.dict()
                except AttributeError:
                    # Если ни один метод не работает, преобразуем объект в словарь вручную
                    activity_dict = {k: v for k, v in activity_data.__dict__.items()
                                    if not k.startswith('_')}

            activity = Activity(
                **activity_dict,
                lesson_id=lesson.id
            )
            activities.append(activity)

        # Only insert activities if there are any
        if activities:
            await self.batch_processor.bulk_insert(activities)

        # Reload the lesson with activities to avoid DetachedInstanceError
        from sqlalchemy import select
        from sqlalchemy.orm import selectinload

        query = (
            select(Lesson)
            .where(Lesson.id == lesson.id)
            .options(selectinload(Lesson.activities))
        )

        result = await self.session.execute(query)
        refreshed_lesson = result.scalar_one()

        return refreshed_lesson

    def _validate_course_data(self, course_data: CourseCreate):
        """Проверяет данные курса на корректность"""
        if not course_data.lessons:
            raise ValidationError("Course must have at least one lesson")

        if not 1 <= len(course_data.lessons) <= 50:
            raise ValidationError("Course must have between 1 and 50 lessons")

        total_duration = sum(
            sum(activity.duration for activity in lesson.activities)
            for lesson in course_data.lessons
        )

        if not 60 <= total_duration <= 10000:
            raise ValidationError("Total course duration must be between 1 and 100 hours")

    @memory_optimized()
    async def generate_course_structure(self, course_data: CourseCreate, user_id: int) -> Course:
        """Генерирует структуру курса с оптимизацией памяти"""
        try:
            # Calculate initial total_duration if lessons are provided
            initial_total_duration = 0
            if course_data.lessons:
                for lesson_data in course_data.lessons:
                    lesson_duration = lesson_data.duration
                    # Add durations of activities if they exist
                    if lesson_data.activities:
                        activity_duration = sum(activity.duration for activity in lesson_data.activities)
                        if activity_duration > 0:
                            lesson_duration = activity_duration
                    initial_total_duration += lesson_duration

            # Создаем новый курс - extract fields from course_data
            course_dict = {
                "name": course_data.name,
                "language": course_data.language,
                "level": course_data.level,
                "start_level": course_data.start_level,
                "target_audience": course_data.target_audience,
                "format": course_data.format,
                "description": course_data.description,
                "exam_prep": course_data.exam_prep,
                "prerequisites": course_data.prerequisites,
                "learning_outcomes": course_data.learning_outcomes,
            }

            course = Course(
                **course_dict,
                total_duration=initial_total_duration,
                creator_id=user_id,

                # Дополнительные поля из расширенной формы
                methodology=course_data.methodology,
                lessons_count=course_data.lessons_count,
                lesson_duration=course_data.lesson_duration,
                main_topics=course_data.main_topics,
                grammar_focus=course_data.grammar_focus,
                vocabulary_focus=course_data.vocabulary_focus,
                student_age=course_data.student_age,
                student_interests=course_data.student_interests,
                student_goals=course_data.student_goals,
                common_mistakes=course_data.common_mistakes,
                include_speaking=course_data.include_speaking,
                include_listening=course_data.include_listening,
                include_reading=course_data.include_reading,
                include_writing=course_data.include_writing,
                include_games=course_data.include_games,
                custom_exam=course_data.custom_exam,
                exam_prep_lessons=course_data.exam_prep_lessons
            )

            self.session.add(course)
            await self.session.flush()

            # Если уже заданы уроки, то создаем их
            if course_data.lessons:
                lessons_data = []
                total_duration = 0
                for lesson_data in course_data.lessons:
                    lesson = await self._create_lesson_with_batch(lesson_data, course.id)
                    lessons_data.append(lesson)
                    total_duration += lesson.duration

                course.total_duration = total_duration
                await self.session.commit()

                # Решение проблемы DetachedInstanceError
                from sqlalchemy import select
                from sqlalchemy.orm import selectinload

                # Создаем запрос с явной загрузкой всех отношений
                query = (
                    select(Course)
                    .where(Course.id == course.id)
                    .options(
                        selectinload(Course.lessons).selectinload(Lesson.activities)
                    )
                )

                # Выполняем запрос
                result = await self.session.execute(query)
                return result.scalar_one()
            else:
                # Генерируем структуру через AI
                generated_structure = await self._generate_with_ai(course_data)

                # Создаем уроки и активности
                lessons_data = []
                total_duration = 0

                for lesson_index, lesson_structure in enumerate(generated_structure['lessons']):
                    # Готовим данные для урока
                    lesson = Lesson(
                        course_id=course.id,
                        title=lesson_structure.get('title', f'Урок {lesson_index+1}'),
                        order=lesson_index+1,
                        objectives=lesson_structure.get('objectives', []),
                        grammar=lesson_structure.get('grammar', []),
                        vocabulary=lesson_structure.get('vocabulary', []),
                        materials=lesson_structure.get('materials', []),
                        homework=lesson_structure.get('homework', {}),
                        duration=lesson_structure.get('duration', 60),
                        is_completed=False
                    )

                    self.session.add(lesson)
                    await self.session.flush()

                    # Создаем активности для урока
                    lesson_duration = 0
                    activities = []

                    for activity_index, activity_data in enumerate(lesson_structure.get('activities', [])):
                        activity = Activity(
                            lesson_id=lesson.id,
                            name=activity_data.get('name', f'Активность {activity_index+1}'),
                            type=activity_data.get('type', 'practice'),
                            duration=activity_data.get('duration', 15),
                            description=activity_data.get('description', ''),
                            materials=activity_data.get('materials', []),
                            objectives=activity_data.get('objectives', [])
                        )
                        activities.append(activity)
                        lesson_duration += activity.duration

                    # Используем batch-processor для вставки активностей
                    await self.batch_processor.bulk_insert(activities)

                    # Обновляем продолжительность урока - УДАЛЕНО, используем значение из lesson_structure.get('duration', 60)
                    # lesson.duration = lesson_duration
                    lessons_data.append(lesson)
                    # Используем исходную длительность урока для общей длительности курса
                    total_duration += lesson.duration # Используем lesson.duration, установленное при создании Lesson

                # Обновляем продолжительность курса
                course.total_duration = total_duration
                await self.session.commit()

                # Решение проблемы DetachedInstanceError: делаем повторный запрос для получения полной структуры курса
                # с предварительно загруженными отношениями
                from sqlalchemy import select
                from sqlalchemy.orm import selectinload

                # Создаем запрос с явной загрузкой всех отношений
                query = (
                    select(Course)
                    .where(Course.id == course.id)
                    .options(
                        selectinload(Course.lessons).selectinload(Lesson.activities)
                    )
                )

                # Выполняем запрос
                result = await self.session.execute(query)
                refreshed_course = result.scalar_one()

                return refreshed_course

        except Exception as e:
            logger.error(f"Error generating course structure: {str(e)}")
            await self.session.rollback()
            raise

    async def _generate_with_ai(self, course_data: CourseCreate) -> Dict[str, Any]:
        """Использует AI для генерации структуры курса"""
        try:
            prompt = self._create_ai_prompt(course_data)
            # --- ДОБАВЛЕНО ЛОГИРОВАНИЕ ---
            import json # Убедимся, что json импортирован
            logger.info(f"Полный промпт для AI: {json.dumps(prompt, ensure_ascii=False, indent=2)}")
            # --- КОНЕЦ ЛОГИРОВАНИЯ ---
            ai_response = await self.ai_service.generate_course_structure(prompt)

            # --- ДОБАВЛЕНО ДЕТАЛЬНОЕ ЛОГИРОВАНИЕ ОТВЕТА AI ---
            logger.info(f"=== ОТЛАДКА: Ответ от AI ===")
            logger.info(f"Тип ответа: {type(ai_response)}")
            logger.info(f"Ключи в ответе: {list(ai_response.keys()) if isinstance(ai_response, dict) else 'Не словарь'}")
            if isinstance(ai_response, dict):
                logger.info(f"Содержит 'lessons': {'lessons' in ai_response}")
                if 'lessons' in ai_response:
                    logger.info(f"Количество уроков: {len(ai_response['lessons']) if isinstance(ai_response['lessons'], list) else 'Не список'}")
                else:
                    logger.info(f"Доступные поля: {list(ai_response.keys())}")
                # Логируем первые 500 символов структуры для анализа
                logger.info(f"Первые 500 символов ответа: {str(ai_response)[:500]}")
            # --- КОНЕЦ ДЕТАЛЬНОГО ЛОГИРОВАНИЯ ---

            self._validate_generated_structure(ai_response)
            return ai_response
        except Exception as e:
            logger.error(f"Error in AI generation: {str(e)}")
            raise

    def _create_ai_prompt(self, course_data: CourseCreate) -> Dict[str, Any]:
        """Создает промпт для генерации полного курса с множеством уроков"""
        # Извлекаем необходимые данные из course_data
        skills = []
        if hasattr(course_data, 'include_speaking') and course_data.include_speaking:
            skills.append("speaking")
        if hasattr(course_data, 'include_listening') and course_data.include_listening:
            skills.append("listening")
        if hasattr(course_data, 'include_reading') and course_data.include_reading:
            skills.append("reading")
        if hasattr(course_data, 'include_writing') and course_data.include_writing:
            skills.append("writing")

        # Определяем общее количество уроков в курсе
        lessons_count = getattr(course_data, 'lessons_count', 12)

        # Получаем методику обучения
        methodology = getattr(course_data, 'methodology', 'communicative')

        # Подготовка к экзамену
        exam_prep = course_data.exam_prep or ""
        custom_exam = getattr(course_data, 'custom_exam', "")
        exam_prep_lessons = getattr(course_data, 'exam_prep_lessons', 0)

        # Получаем начальный уровень (если указан) или используем целевой уровень
        start_level = course_data.start_level.value if hasattr(course_data.start_level, 'value') and course_data.start_level else (
            course_data.start_level if course_data.start_level else course_data.level.value if hasattr(course_data.level, 'value') else course_data.level
        )

        # Информация о студенте
        student_age = getattr(course_data, 'student_age', "")
        student_interests = getattr(course_data, 'student_interests', "")
        student_goals = getattr(course_data, 'student_goals', "")
        common_mistakes = getattr(course_data, 'common_mistakes', "")

        # Создаем контекст для промпта
        prompt_context = {
            "course_name": course_data.name,
            "language": course_data.language,
            "level": course_data.level.value if hasattr(course_data.level, 'value') else course_data.level,
            "start_level": start_level,
            "target_audience": course_data.target_audience.value if hasattr(course_data.target_audience, 'value') else course_data.target_audience,
            "format": course_data.format.value if hasattr(course_data.format, 'value') else course_data.format,
            "methodology": methodology,
            "description": course_data.description,
            "exam_prep": exam_prep,
            "custom_exam": custom_exam,
            "exam_prep_lessons": exam_prep_lessons,
            "prerequisites": course_data.prerequisites,
            "learning_outcomes": course_data.learning_outcomes,
            "lessons_count": lessons_count,
            "lesson_duration": getattr(course_data, 'lesson_duration', 60),
            "main_topics": getattr(course_data, 'main_topics', ""),
            "grammar_focus": getattr(course_data, 'grammar_focus', ""),
            "vocabulary_focus": getattr(course_data, 'vocabulary_focus', ""),
            "skills": skills,
            "include_games": getattr(course_data, 'include_games', True),
            "student_age": student_age,
            "student_interests": student_interests,
            "student_goals": student_goals,
            "common_mistakes": common_mistakes,
        }

        # Описания методик обучения
        methodology_descriptions = {
            'communicative': "Коммуникативная методика фокусируется на интерактивном общении и развитии разговорных навыков через реальные ситуации общения",
            'task-based': "Проектно-ориентированное обучение строится на выполнении практических задач, которые требуют использования языка в реальных ситуациях",
            'natural': "Натуральный метод имитирует естественное усвоение языка через погружение в языковую среду без явного изучения грамматики",
            'lexical': "Лексический подход концентрируется на изучении лексических блоков, устойчивых выражений и фраз вместо отдельных слов",
            'grammar-translation': "Грамматико-переводной метод основан на систематическом изучении грамматических правил и переводе текстов",
            'audio-lingual': "Аудиолингвальный метод основан на прослушивании и повторении, с акцентом на формировании языковых привычек через повторение образцов",
            'direct': "Прямой метод исключает использование родного языка и перевода, фокусируясь на полном погружении в изучаемый язык",
            'total-physical-response': "Метод физического реагирования связывает язык с физическими действиями, особенно эффективен для начинающих и детей"
        }

        # Получаем описание выбранной методики
        methodology_description = methodology_descriptions.get(methodology, "")

        # Подготовка описания экзамена
        exam_description = ""
        if exam_prep:
            if exam_prep_lessons > 0:
                exam_description = f"\nПодготовка к экзамену {exam_prep} {custom_exam if custom_exam else ''} должна быть включена в последние {exam_prep_lessons} уроков курса."
            else:
                exam_description = f"\nВключите элементы подготовки к экзамену {exam_prep} {custom_exam if custom_exam else ''} во все уроки курса."

        # Подготовка описания целей и ошибок студента
        student_info = ""
        if student_age:
            student_info += f"\nВозраст студента: {student_age}."
        if student_interests:
            student_info += f"\nИнтересы студента: {student_interests}."
        if student_goals:
            student_info += f"\nЦели студента: {student_goals}."
        if common_mistakes:
            student_info += f"\nТипичные ошибки студента: {common_mistakes}."

        # Подготовка описания уровней
        level_transition = ""
        if start_level and start_level != prompt_context["level"]:
            level_transition = f"\nКурс должен обеспечить плавный переход с уровня {start_level} на уровень {prompt_context['level']}."

        prompt = {
            "task": "generate_full_course",
            "context": prompt_context,
            "requirements": f"""
            Создайте полный курс по изучению языка {course_data.language} для {prompt_context['target_audience']} уровня {prompt_context['level']}.
            Курс должен содержать {lessons_count} полных занятий, каждое продолжительностью {prompt_context['lesson_duration']} минут.

            {level_transition}

            {student_info}

            Используйте методику обучения: {methodology} ({methodology_description})

            {exam_description}

            Каждый урок должен включать:
            1. Название урока
            2. Цели обучения
            3. Лексику к уроку
            4. Грамматику к уроку
            5. Структурированные активности с указанием продолжительности
            6. Материалы для занятия
            7. Домашнее задание

            Важно:
            - Уроки должны быть последовательными и строиться друг на друге
            - Учитывайте уровень сложности {prompt_context['level']}
            - Адаптируйте содержание для {prompt_context['target_audience']}
            - Используйте формат обучения: {prompt_context['format']}
            - Содержание должно соответствовать выбранной методике
            - Учитывайте интересы и возраст студента при создании активностей
            - Обратите особое внимание на исправление типичных ошибок студента
            - Помогайте достигать указанных целей обучения

            ОБЯЗАТЕЛЬНО верните ответ в следующем JSON формате:
            {{
                "title": "Название курса",
                "description": "Описание курса",
                "level": "{prompt_context['level']}",
                "language": "{course_data.language}",
                "total_lessons": {lessons_count},
                "lessons": [
                    {{
                        "title": "Название урока 1",
                        "lesson_number": 1,
                        "objectives": ["Цель 1", "Цель 2", "Цель 3"],
                        "vocabulary": ["слово1", "слово2", "фраза1"],
                        "grammar": ["Грамматическая тема 1", "Правило 1"],
                        "activities": [
                            {{
                                "name": "Название активности",
                                "duration": 15,
                                "description": "Описание активности"
                            }}
                        ],
                        "materials": ["Материал 1", "Материал 2"],
                        "homework": {{
                            "description": "Описание домашнего задания",
                            "tasks": ["Задание 1", "Задание 2"]
                        }},
                        "duration": {prompt_context['lesson_duration']}
                    }}
                ]
            }}

            ВАЖНО:
            - Ответ должен быть ТОЛЬКО валидным JSON
            - Обязательно включите поле "lessons" с массивом уроков
            - Каждый урок должен содержать все указанные поля
            - Создайте ровно {lessons_count} уроков
            """,
            "output_format": "json"
        }

        return prompt

    def _validate_generated_structure(self, structure: Dict[str, Any]) -> None:
        """Проверяет корректность сгенерированной структуры курса"""
        logger.info("=== НАЧАЛО ВАЛИДАЦИИ СТРУКТУРЫ КУРСА ===")

        if not structure:
            logger.error("Структура курса пустая")
            raise ValidationError("Generated structure is empty")

        logger.info(f"Доступные ключи в структуре: {list(structure.keys())}")

        if "lessons" not in structure:
            logger.error("В структуре отсутствует поле 'lessons'")
            logger.error(f"Доступные поля: {list(structure.keys())}")
            raise ValidationError("Generated structure missing lessons")

        lessons = structure["lessons"]
        if not lessons:
            logger.error("Поле 'lessons' пустое")
            raise ValidationError("Generated structure has no lessons")

        if not isinstance(lessons, list):
            logger.error(f"Поле 'lessons' не является списком, тип: {type(lessons)}")
            raise ValidationError("Generated structure lessons is not a list")

        logger.info(f"Найдено {len(lessons)} уроков для валидации")

        for i, lesson in enumerate(lessons):
            logger.debug(f"Валидация урока {i+1}")

            if not isinstance(lesson, dict):
                logger.error(f"Урок {i+1} не является словарем, тип: {type(lesson)}")
                raise ValidationError(f"Lesson {i+1} is not a dictionary")

            # Проверяем обязательные поля
            required_fields = ["title", "objectives", "activities"]
            for field in required_fields:
                if field not in lesson:
                    logger.error(f"Урок {i+1} не содержит поле '{field}'. Доступные поля: {list(lesson.keys())}")
                    raise ValidationError(f"Lesson {i+1} is missing {field}")

            # Дополнительные проверки
            if not lesson["title"] or not isinstance(lesson["title"], str):
                logger.error(f"Урок {i+1} имеет некорректное название: {lesson.get('title')}")
                raise ValidationError(f"Lesson {i+1} has invalid title")

            if not isinstance(lesson["objectives"], list) or not lesson["objectives"]:
                logger.error(f"Урок {i+1} имеет некорректные цели: {lesson.get('objectives')}")
                raise ValidationError(f"Lesson {i+1} has invalid objectives")

            if not isinstance(lesson["activities"], list) or not lesson["activities"]:
                logger.error(f"Урок {i+1} имеет некорректные активности: {lesson.get('activities')}")
                raise ValidationError(f"Lesson {i+1} has invalid activities")

        logger.info(f"✅ Валидация успешно завершена для {len(lessons)} уроков")
        logger.info("=== КОНЕЦ ВАЛИДАЦИИ СТРУКТУРЫ КУРСА ===")

    @memory_optimized()
    async def generate_lesson_materials(self, lesson_id: int, types: List[str]) -> Dict[str, Any]:
        """
        Генерирует запрошенные материалы для урока и возвращает их в виде словаря.
        """
        logger.info(f"Запрос на генерацию материалов {types} для урока {lesson_id}")
        try:
            # Проверяем кэш (можно оставить или убрать, если генерация по запросу не кэшируется)
            # cache_key = f"lesson_materials:{lesson_id}:{'-'.join(sorted(types))}" # Сортируем типы для консистентности ключа
            # cached_materials = await self.cache_service.get_cached_data(cache_key)
            # if cached_materials:
            #     logger.info(f"Возвращаем кэшированные материалы для урока {lesson_id}, типы: {types}")
            #     return cached_materials

            # Импортируем selectinload здесь, если он нужен
            from sqlalchemy.orm import selectinload
            query = await self.query_optimizer.optimize_query(
                select(Lesson)
                .where(Lesson.id == lesson_id)
                .options(selectinload(Lesson.course))  # Загружаем курс для контекста
            )
            result = await self.session.execute(query)
            lesson = result.scalar_one_or_none()

            if not lesson:
                raise NotFoundException(f"Урок с ID {lesson_id} не найден")

            # Словарь для хранения результатов
            materials_response = {}
            generation_tasks = []
            requested_material_types = []  # Список для хранения типов, которые будем генерировать

            # Формируем задачи для асинхронной генерации только запрошенных типов
            valid_types = ['exercises', 'presentation', 'handouts', 'assessments']  # Определяем валидные типы
            for material_type in types:
                material_type_lower = material_type.lower()
                if material_type_lower not in valid_types:
                    logger.warning(f"Неизвестный тип материала запрошен: {material_type}")
                    continue  # Пропускаем неизвестные типы

                requested_material_types.append(material_type_lower)  # Добавляем валидный тип

                if material_type_lower == 'exercises':
                    generation_tasks.append(self._generate_exercises(lesson))
                elif material_type_lower == 'presentation':
                    generation_tasks.append(self._generate_presentations(lesson))
                elif material_type_lower == 'handouts':
                    generation_tasks.append(self._generate_handouts(lesson))
                elif material_type_lower == 'assessments':
                    generation_tasks.append(self._generate_assessments(lesson))
                # Добавить другие типы при необходимости

            if not generation_tasks:
                logger.warning(f"Не запрошено ни одного известного типа материалов для урока {lesson_id}")
                return {}  # Возвращаем пустой словарь

            # Асинхронно выполняем все задачи генерации
            logger.info(f"Запуск генерации материалов: {requested_material_types} для урока {lesson_id}")
            results = await asyncio.gather(*generation_tasks, return_exceptions=True)  # Собираем исключения

            # Объединяем результаты в словарь ответа
            for material_type, result in zip(requested_material_types, results):
                if isinstance(result, Exception):
                    logger.error(
                        f"Ошибка при генерации материала типа '{material_type}' для урока {lesson_id}: {result}")
                    materials_response[material_type] = {"error": f"Failed to generate {material_type}: {str(result)}"}
                else:
                    materials_response[material_type] = result

            # Не обновляем lesson.materials и не кэшируем здесь
            # await self.cache_service.cache_data(cache_key, materials_response, ttl=3600) # Можно добавить кэширование, если нужно

            logger.info(f"Материалы {list(materials_response.keys())} сгенерированы для урока {lesson_id}")
            return materials_response  # Возвращаем словарь с результатами

        except NotFoundException:
            raise
        except Exception as e:
            logger.error(f"Error generating materials for lesson {lesson_id}: {str(e)}")
            # Не откатываем транзакцию, т.к. мы только читали данные урока
            # await self.session.rollback()
            raise

    async def _generate_exercises(self, lesson: Lesson) -> List[Dict[str, Any]]:
        """Генерирует упражнения для урока"""
        prompt = self._create_exercises_prompt(lesson)
        exercises = await self.ai_service.generate_exercises(prompt)
        return exercises

    async def _generate_presentations(self, lesson: Lesson) -> Dict[str, Any]:
        """Генерирует презентации для урока с оптимизированной AI-генерацией"""
        try:
            # Проверяем кэш
            cache_key = f"presentation:{lesson.id}"
            cached_presentation = await self.cache_service.get_cached_data(cache_key)
            if cached_presentation:
                return cached_presentation

            # Формируем промпт для AI
            prompt = self._create_presentation_prompt(lesson)

            # Генерируем слайды
            slides_content = await self.ai_service.generate_content(prompt)

            # Структурируем презентацию
            presentation = {
                "title": lesson.title,
                "theme": "default",
                "slides": [
                    {
                        "type": "title",
                        "content": {
                            "title": lesson.title,
                            "subtitle": f"Level: {lesson.course.level}"
                        }
                    }
                ],
                "style": {
                    "colors": {
                        "primary": "#4CAF50",
                        "secondary": "#2196F3",
                        "accent": "#FF9800"
                    },
                    "fonts": {
                        "heading": "Roboto",
                        "body": "Open Sans"
                    }
                },
                "metadata": {
                    "created_at": datetime.now(timezone.utc).isoformat(),
                    "lesson_id": lesson.id,
                    "duration": lesson.duration,
                    "objectives": lesson.objectives
                }
            }

            # Обрабатываем и структурируем сгенерированный контент
            slide_sections = slides_content.split("---")
            for section in slide_sections:
                if not section.strip():
                    continue

                # Определяем тип слайда на основе содержимого
                slide_type = "content"
                if "exercise" in section.lower():
                    slide_type = "exercise"
                elif "summary" in section.lower():
                    slide_type = "summary"
                elif "quiz" in section.lower():
                    slide_type = "quiz"

                slide = {
                    "type": slide_type,
                    "content": self._process_slide_content(section),
                    "notes": self._generate_speaker_notes(section)
                }
                presentation["slides"].append(slide)

            # Добавляем финальный слайд
            presentation["slides"].append({
                "type": "summary",
                "content": {
                    "title": "Thank you!",
                    "points": lesson.objectives
                }
            })

            # Кэшируем результат
            await self.cache_service.cache_data(cache_key, presentation, ttl=3600)

            return presentation

        except Exception as e:
            logger.error(f"Error generating presentation: {str(e)}")
            raise

    async def _generate_handouts(self, lesson: Lesson) -> List[Dict[str, Any]]:
        """Генерирует раздаточные материалы для урока"""
        try:
            # Проверяем кэш
            cache_key = f"handouts:{lesson.id}"
            cached_handouts = await self.cache_service.get_cached_data(cache_key)
            if cached_handouts:
                return cached_handouts

            # Генерируем различные типы материалов
            handouts = []

            # 1. Основной конспект
            summary_prompt = self._create_summary_prompt(lesson)
            summary_content = await self.ai_service.generate_content(summary_prompt)
            handouts.append({
                "type": "summary",
                "title": f"{lesson.title} - Summary",
                "content": summary_content,
                "format": "markdown",
                "pages": self._estimate_pages(summary_content)
            })

            # 2. Рабочие листы
            worksheet_prompt = self._create_worksheet_prompt(lesson)
            worksheet_content = await self.ai_service.generate_content(worksheet_prompt)
            handouts.append({
                "type": "worksheet",
                "title": f"{lesson.title} - Worksheet",
                "content": worksheet_content,
                "format": "markdown",
                "pages": self._estimate_pages(worksheet_content)
            })

            # 3. Словарь/глоссарий
            if lesson.vocabulary:
                vocabulary_content = await self._generate_vocabulary_handout(lesson)
                handouts.append({
                    "type": "vocabulary",
                    "title": f"{lesson.title} - Vocabulary List",
                    "content": vocabulary_content,
                    "format": "markdown",
                    "pages": self._estimate_pages(vocabulary_content)
                })

            # 4. Грамматические карточки
            if lesson.grammar:
                grammar_content = await self._generate_grammar_handout(lesson)
                handouts.append({
                    "type": "grammar",
                    "title": f"{lesson.title} - Grammar Reference",
                    "content": grammar_content,
                    "format": "markdown",
                    "pages": self._estimate_pages(grammar_content)
                })

            # Добавляем метаданные ко всем материалам
            for handout in handouts:
                handout.update({
                    "metadata": {
                        "lesson_id": lesson.id,
                        "course_id": lesson.course_id,
                        "level": lesson.course.level,
                        "created_at": datetime.now(timezone.utc).isoformat(),
                        "language": lesson.course.language
                    }
                })

            # Кэшируем результат
            await self.cache_service.cache_data(cache_key, handouts, ttl=3600)

            return handouts

        except Exception as e:
            logger.error(f"Error generating handouts: {str(e)}")
            raise

    async def _generate_assessments(self, lesson: Lesson) -> Dict[str, Any]:
        """Генерирует материалы для оценки знаний"""
        try:
            # Проверяем кэш
            cache_key = f"assessments:{lesson.id}"
            cached_assessments = await self.cache_service.get_cached_data(cache_key)
            if cached_assessments:
                return cached_assessments

            # Формируем набор оценочных материалов
            assessments = {
                "pre_assessment": await self._generate_pre_assessment(lesson),
                "formative_assessment": await self._generate_formative_assessment(lesson),
                "summative_assessment": await self._generate_summative_assessment(lesson),
                "self_assessment": await self._generate_self_assessment(lesson),
                "rubrics": await self._generate_assessment_rubrics(lesson),
                "metadata": {
                    "lesson_id": lesson.id,
                    "course_id": lesson.course_id,
                    "level": lesson.course.level,
                    "learning_objectives": lesson.objectives,
                    "created_at": datetime.now(timezone.utc).isoformat()
                }
            }

            # Добавляем общую статистику
            assessments["statistics"] = {
                "total_questions": sum(
                    len(assessment.get("questions", []))
                    for assessment in assessments.values()
                    if isinstance(assessment, dict)
                ),
                "difficulty_distribution": {
                    "easy": 30,
                    "medium": 40,
                    "hard": 30
                },
                "estimated_duration": {
                    "pre_assessment": 15,
                    "formative_assessment": 20,
                    "summative_assessment": 30,
                    "self_assessment": 10
                }
            }

            # Кэшируем результат
            await self.cache_service.cache_data(cache_key, assessments, ttl=3600)

            return assessments

        except Exception as e:
            logger.error(f"Error generating assessments: {str(e)}")
            raise

    async def _generate_pre_assessment(self, lesson: Lesson) -> Dict[str, Any]:
        """Генерирует предварительную оценку"""
        prompt = self._create_pre_assessment_prompt(lesson)
        content = await self.ai_service.generate_content(prompt)
        return self._structure_assessment(content, "pre_assessment")

    async def _generate_formative_assessment(self, lesson: Lesson) -> Dict[str, Any]:
        """Генерирует формирующую оценку"""
        prompt = self._create_formative_assessment_prompt(lesson)
        content = await self.ai_service.generate_content(prompt)
        return self._structure_assessment(content, "formative_assessment")

    async def _generate_summative_assessment(self, lesson: Lesson) -> Dict[str, Any]:
        """Генерирует итоговую оценку"""
        prompt = self._create_summative_assessment_prompt(lesson)
        content = await self.ai_service.generate_content(prompt)
        return self._structure_assessment(content, "summative_assessment")

    async def _generate_self_assessment(self, lesson: Lesson) -> Dict[str, Any]:
        """Генерирует материалы для самооценки"""
        prompt = self._create_self_assessment_prompt(lesson)
        content = await self.ai_service.generate_content(prompt)
        return self._structure_assessment(content, "self_assessment")

    async def _generate_assessment_rubrics(self, lesson: Lesson) -> Dict[str, Any]:
        """Генерирует рубрики оценивания"""
        prompt = self._create_rubrics_prompt(lesson)
        content = await self.ai_service.generate_content(prompt)
        return self._structure_rubrics(content)

    def _structure_assessment(self, content: str, assessment_type: str) -> Dict[str, Any]:
        """Структурирует оценочные материалы"""
        # Преобразуем сгенерированный контент в структурированный формат
        return {
            "type": assessment_type,
            "questions": self._parse_questions(content),
            "scoring_guide": self._generate_scoring_guide(content),
            "time_limit": self._get_time_limit(assessment_type),
            "passing_score": self._get_passing_score(assessment_type)
        }

    def _structure_rubrics(self, content: str) -> Dict[str, Any]:
        """Структурирует рубрики оценивания"""
        return {
            "criteria": self._parse_rubric_criteria(content),
            "scoring_levels": self._parse_scoring_levels(content),
            "usage_guide": self._generate_rubric_guide(content)
        }

    @memory_optimized()
    async def clone_course(self, course_id: int, new_name: str, user_id: int) -> Course:
        """Клонирует существующий курс с оптимизацией"""
        try:
            # Получаем исходный курс
            source_course = await self.get_course(course_id)
            if not source_course:
                raise NotFoundException("Course not found")

            # Создаем новый курс
            new_course = Course(
                name=new_name,
                language=source_course.language,
                level=source_course.level,
                target_audience=source_course.target_audience,
                format=source_course.format,
                description=source_course.description,
                exam_prep=source_course.exam_prep,
                prerequisites=source_course.prerequisites.copy(),
                learning_outcomes=source_course.learning_outcomes.copy(),
                creator_id=user_id,
                is_used=False,
                usage_count=0
            )
            self.session.add(new_course)
            await self.session.flush()

            # Получаем все уроки исходного курса оптимизированным запросом
            query = await self.query_optimizer.optimize_query(
                select(Lesson).where(Lesson.course_id == course_id)
            )
            result = await self.session.execute(query)
            source_lessons = result.scalars().all()

            # Подготавливаем данные для батч-вставки
            new_lessons = []
            activities_data = []

            for lesson in source_lessons:
                new_lesson = Lesson(
                    course_id=new_course.id,
                    title=lesson.title,
                    duration=lesson.duration,
                    order=lesson.order,
                    objectives=lesson.objectives.copy(),
                    grammar=lesson.grammar.copy(),
                    vocabulary=lesson.vocabulary.copy(),
                    materials=lesson.materials.copy(),
                    homework=lesson.homework.copy()
                )
                self.session.add(new_lesson)
                await self.session.flush()
                new_lessons.append(new_lesson)

                # Клонируем активности
                activity_query = await self.query_optimizer.optimize_query(
                    select(Activity).where(Activity.lesson_id == lesson.id)
                )
                activity_result = await self.session.execute(activity_query)
                activities = activity_result.scalars().all()

                for activity in activities:
                    new_activity = Activity(
                        lesson_id=new_lesson.id,
                        name=activity.name,
                        type=activity.type,
                        duration=activity.duration,
                        description=activity.description,
                        materials=activity.materials.copy(),
                        objectives=activity.objectives.copy()
                    )
                    activities_data.append(new_activity)

            # Батч-вставка всех активностей
            await self.batch_processor.bulk_insert(activities_data)

            await self.session.commit()
            await self.session.refresh(new_course)

            return new_course

        except Exception as e:
            logger.error(f"Error cloning course: {str(e)}")
            await self.session.rollback()
            raise

    @memory_optimized()
    async def export_course(self, course_id: int, format: str) -> Tuple[bytes, str]:
        """
        Экспортирует курс в выбранном формате (PDF или DOCX) и возвращает байты файла и имя.
        """
        logger.info(f"Начало экспорта курса {course_id} в формат {format}")
        try:
            # Импортируем selectinload здесь
            from sqlalchemy.orm import selectinload
            # Получаем курс с полной структурой уроков и активностей
            query = await self.query_optimizer.optimize_query(
                select(Course)
                .where(Course.id == course_id)
                .options(
                    selectinload(Course.lessons)
                    .selectinload(Lesson.activities)  # Загружаем активности сразу
                )
            )
            result = await self.session.execute(query)
            course = result.scalar_one_or_none()

            if not course:
                raise NotFoundException(f"Курс с ID {course_id} не найден")

            # Сортируем уроки по order на всякий случай
            course.lessons.sort(key=lambda l: l.order)

            # Используем только ID курса в имени файла, чтобы избежать проблем с кодировкой
            filename = f"course_{course_id}.{format}"
            file_content_bytes: bytes

            if format == 'pdf':
                # --- Логика генерации PDF ---
                # Убедитесь, что reportlab установлен: pip install reportlab
                try:
                    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
                    from reportlab.lib.styles import getSampleStyleSheet
                    from reportlab.lib.units import inch
                    from reportlab.pdfbase import pdfmetrics
                    from reportlab.pdfbase.ttfonts import TTFont
                    import io
                    import os

                    # Регистрируем шрифты с поддержкой кириллицы
                    # Используем DejaVuSans, который поддерживает кириллицу
                    font_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', '..', 'static', 'fonts')

                    # Проверяем наличие шрифтов в папке static/fonts
                    dejavu_path = os.path.join(font_path, 'DejaVuSans.ttf')
                    dejavu_bold_path = os.path.join(font_path, 'DejaVuSans-Bold.ttf')

                    fonts_registered = False

                    if os.path.exists(dejavu_path):
                        pdfmetrics.registerFont(TTFont('DejaVuSans', dejavu_path))
                        logger.info(f"Зарегистрирован шрифт DejaVuSans из {dejavu_path}")
                        fonts_registered = True

                    if os.path.exists(dejavu_bold_path):
                        pdfmetrics.registerFont(TTFont('DejaVuSans-Bold', dejavu_bold_path))
                        logger.info(f"Зарегистрирован шрифт DejaVuSans-Bold из {dejavu_bold_path}")
                        fonts_registered = True

                    # Регистрируем шрифтовые пары для поддержки кириллицы
                    if fonts_registered:
                        from reportlab.pdfbase.pdfmetrics import registerFontFamily
                        registerFontFamily('DejaVuSans', normal='DejaVuSans', bold='DejaVuSans-Bold')
                        logger.info("Зарегистрировано семейство шрифтов DejaVuSans")

                    if not fonts_registered:
                        # Если шрифтов нет в папке, используем стандартный Helvetica
                        logger.warning(f"Шрифты DejaVuSans не найдены в {font_path}, используем стандартные шрифты")

                    # Создаем буфер для PDF
                    buffer = io.BytesIO()

                    # Создаем документ с указанием кодировки
                    doc = SimpleDocTemplate(
                        buffer,
                        pagesize=(8.5 * inch, 11 * inch),
                        leftMargin=0.75 * inch,
                        rightMargin=0.75 * inch,
                        topMargin=1 * inch,
                        bottomMargin=1 * inch,
                        encoding='utf-8',  # Явно указываем кодировку UTF-8
                        title=course.name,  # Устанавливаем заголовок документа
                        author='Course Generator',  # Устанавливаем автора
                        subject=f'Course: {course.name}'  # Устанавливаем тему
                    )
                    styles = getSampleStyleSheet()

                    # Создаем копии стилей с поддержкой кириллицы
                    if fonts_registered:
                        # Импортируем необходимые классы для создания стилей
                        from reportlab.lib.styles import ParagraphStyle
                        from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
                        from reportlab.lib import colors

                        # Создаем собственные стили с поддержкой кириллицы
                        styles.add(ParagraphStyle(
                            name='CustomTitle',
                            fontName='DejaVuSans-Bold',
                            fontSize=18,
                            leading=22,
                            alignment=TA_CENTER,
                            spaceAfter=12,
                            textColor=colors.black
                        ))

                        styles.add(ParagraphStyle(
                            name='CustomHeading1',
                            fontName='DejaVuSans-Bold',
                            fontSize=16,
                            leading=20,
                            alignment=TA_LEFT,
                            spaceAfter=10,
                            textColor=colors.black
                        ))

                        styles.add(ParagraphStyle(
                            name='CustomHeading2',
                            fontName='DejaVuSans-Bold',
                            fontSize=14,
                            leading=18,
                            alignment=TA_LEFT,
                            spaceAfter=8,
                            textColor=colors.black
                        ))

                        styles.add(ParagraphStyle(
                            name='CustomHeading3',
                            fontName='DejaVuSans-Bold',
                            fontSize=12,
                            leading=16,
                            alignment=TA_LEFT,
                            spaceAfter=6,
                            textColor=colors.black
                        ))

                        styles.add(ParagraphStyle(
                            name='CustomNormal',
                            fontName='DejaVuSans',
                            fontSize=10,
                            leading=14,
                            alignment=TA_LEFT,
                            spaceAfter=6,
                            textColor=colors.black
                        ))

                        # Применяем шрифты к стандартным стилям
                        for style_name in styles.byName:
                            style = styles[style_name]
                            if hasattr(style, 'fontName'):
                                if style_name in ['h1', 'h2', 'h3', 'h4', 'Heading1', 'Heading2', 'Heading3', 'Heading4', 'Title']:
                                    # Для заголовков используем жирный шрифт
                                    if 'DejaVuSans-Bold' in pdfmetrics.getRegisteredFontNames():
                                        style.fontName = 'DejaVuSans-Bold'
                                else:
                                    # Для обычного текста используем обычный шрифт
                                    if 'DejaVuSans' in pdfmetrics.getRegisteredFontNames():
                                        style.fontName = 'DejaVuSans'

                    story = []

                    # Титульная страница
                    story.append(Paragraph(course.name, styles['CustomTitle'] if fonts_registered else styles['h1']))
                    story.append(Spacer(1, 0.2 * inch))
                    story.append(Paragraph(f"Язык: {course.language}, Уровень: {course.level}",
                                          styles['CustomHeading3'] if fonts_registered else styles['h3']))
                    story.append(
                        Paragraph(f"Аудитория: {course.target_audience}, Формат: {course.format}",
                                 styles['CustomHeading3'] if fonts_registered else styles['h3']))
                    if course.description:
                        story.append(Spacer(1, 0.1 * inch))
                        story.append(Paragraph(f"Описание: {course.description}",
                                             styles['CustomNormal'] if fonts_registered else styles['Normal']))
                    story.append(PageBreak())

                    # Уроки
                    for i, lesson in enumerate(course.lessons):
                        story.append(Paragraph(f"Урок {i + 1}: {lesson.title}",
                                             styles['CustomHeading1'] if fonts_registered else styles['h2']))
                        story.append(Spacer(1, 0.1 * inch))
                        story.append(Paragraph(f"Продолжительность: {lesson.duration} минут",
                                             styles['CustomNormal'] if fonts_registered else styles['Italic']))
                        story.append(Spacer(1, 0.1 * inch))

                        story.append(Paragraph("Цели:",
                                             styles['CustomHeading3'] if fonts_registered else styles['h3']))
                        for obj in lesson.objectives:
                            story.append(Paragraph(f"- {obj}",
                                                 styles['CustomNormal'] if fonts_registered else styles['Normal']))
                        story.append(Spacer(1, 0.1 * inch))

                        story.append(Paragraph("Грамматика:",
                                             styles['CustomHeading3'] if fonts_registered else styles['h3']))
                        for gram in lesson.grammar:
                            story.append(Paragraph(f"- {gram}",
                                                 styles['CustomNormal'] if fonts_registered else styles['Normal']))
                        story.append(Spacer(1, 0.1 * inch))

                        story.append(Paragraph("Лексика:",
                                             styles['CustomHeading3'] if fonts_registered else styles['h3']))
                        for vocab in lesson.vocabulary:
                            story.append(Paragraph(f"- {vocab}",
                                                 styles['CustomNormal'] if fonts_registered else styles['Normal']))
                        story.append(Spacer(1, 0.1 * inch))

                        story.append(Paragraph("Активности:",
                                             styles['CustomHeading3'] if fonts_registered else styles['h3']))
                        for act in lesson.activities:
                            story.append(Paragraph(f"- {act.name} ({act.type}, {act.duration} мин)",
                                                 styles['CustomHeading3'] if fonts_registered else styles['h4']))
                            if act.description:
                                story.append(Paragraph(act.description,
                                                     styles['CustomNormal'] if fonts_registered else styles['Normal']))
                            story.append(Spacer(1, 0.05 * inch))
                        story.append(Spacer(1, 0.1 * inch))

                        if lesson.homework:
                            story.append(Paragraph("Домашнее задание:",
                                                 styles['CustomHeading3'] if fonts_registered else styles['h3']))
                            if isinstance(lesson.homework, dict):
                                if lesson.homework.get('description'):
                                    story.append(Paragraph(lesson.homework['description'],
                                                         styles['CustomNormal'] if fonts_registered else styles['Normal']))
                                if lesson.homework.get('tasks'):
                                    for task in lesson.homework['tasks']:
                                        story.append(Paragraph(f"- {task}",
                                                             styles['CustomNormal'] if fonts_registered else styles['Normal']))
                            else:  # Если homework - строка
                                story.append(Paragraph(str(lesson.homework),
                                                     styles['CustomNormal'] if fonts_registered else styles['Normal']))
                            story.append(Spacer(1, 0.1 * inch))

                        if i < len(course.lessons) - 1:
                            story.append(PageBreak())

                    # Создаем PDF с указанием кодировки и дополнительными метаданными
                    from reportlab.lib.pagesizes import letter

                    # Добавляем метаданные документа
                    doc_info = {
                        'Title': course.name,
                        'Author': 'Course Generator',
                        'Subject': f'Course: {course.name}',
                        'Creator': 'ReportLab PDF Library',
                        'Producer': 'Course Generator System'
                    }

                    # Создаем функцию для первой страницы
                    def first_page(canvas, doc):
                        canvas.saveState()
                        # Устанавливаем метаданные документа
                        canvas.setTitle(doc_info.get('Title', ''))
                        canvas.setAuthor(doc_info.get('Author', ''))
                        canvas.setSubject(doc_info.get('Subject', ''))
                        canvas.setCreator(doc_info.get('Creator', ''))
                        # Устанавливаем кодировку
                        canvas.setPageCompression(1)
                        canvas.restoreState()

                    # Создаем функцию для остальных страниц
                    def later_pages(canvas, doc):
                        canvas.saveState()
                        # Добавляем номер страницы
                        canvas.setFont('DejaVuSans' if fonts_registered else 'Helvetica', 9)
                        canvas.drawString(inch, 0.75 * inch, f"Страница {doc.page}")
                        canvas.restoreState()

                    # Строим документ с указанием функций для страниц
                    doc.build(story, onFirstPage=first_page, onLaterPages=later_pages)

                    file_content_bytes = buffer.getvalue()
                    logger.info(f"PDF для курса {course_id} успешно сгенерирован.")

                except ImportError:
                    logger.error("Библиотека reportlab не установлена. Невозможно экспортировать в PDF.")
                    raise ValidationError("Экспорт в PDF недоступен. Установите reportlab.")
                except Exception as pdf_err:
                    logger.error(f"Ошибка при генерации PDF для курса {course_id}: {pdf_err}")
                    raise

            elif format == 'docx':
                # --- Логика генерации DOCX ---
                # Убедитесь, что python-docx установлен: pip install python-docx
                try:
                    from docx import Document
                    from docx.shared import Inches, Pt
                    import io

                    document = Document()
                    # Титульная страница
                    document.add_heading(course.name, level=1)
                    document.add_paragraph(f"Язык: {course.language}, Уровень: {course.level}")
                    document.add_paragraph(f"Аудитория: {course.target_audience}, Формат: {course.format}")
                    if course.description: document.add_paragraph(f"Описание: {course.description}")
                    document.add_page_break()

                    # Уроки
                    for i, lesson in enumerate(course.lessons):
                        document.add_heading(f"Урок {i + 1}: {lesson.title}", level=2)
                        document.add_paragraph(f"Продолжительность: {lesson.duration} минут").italic = True

                        document.add_heading("Цели:", level=3)
                        for obj in lesson.objectives: document.add_paragraph(obj, style='List Bullet')

                        document.add_heading("Грамматика:", level=3)
                        for gram in lesson.grammar: document.add_paragraph(gram, style='List Bullet')

                        document.add_heading("Лексика:", level=3)
                        for vocab in lesson.vocabulary: document.add_paragraph(vocab, style='List Bullet')

                        document.add_heading("Активности:", level=3)
                        for act in lesson.activities:
                            p = document.add_paragraph()
                            p.add_run(f"{act.name} ({act.type}, {act.duration} мин)").bold = True
                            if act.description: document.add_paragraph(act.description)

                        if lesson.homework:
                            document.add_heading("Домашнее задание:", level=3)
                            if isinstance(lesson.homework, dict):
                                if lesson.homework.get('description'): document.add_paragraph(
                                    lesson.homework['description'])
                                if lesson.homework.get('tasks'):
                                    for task in lesson.homework['tasks']: document.add_paragraph(task,
                                                                                                 style='List Bullet')
                            else:
                                document.add_paragraph(str(lesson.homework))

                        if i < len(course.lessons) - 1:
                            document.add_page_break()

                    # Сохраняем в байтовый поток
                    file_stream = io.BytesIO()
                    document.save(file_stream)
                    file_stream.seek(0)
                    file_content_bytes = file_stream.read()
                    logger.info(f"DOCX для курса {course_id} успешно сгенерирован.")

                except ImportError:
                    logger.error("Библиотека python-docx не установлена. Невозможно экспортировать в DOCX.")
                    raise ValidationError("Экспорт в DOCX недоступен. Установите python-docx.")
                except Exception as docx_err:
                    logger.error(f"Ошибка при генерации DOCX для курса {course_id}: {docx_err}")
                    raise
            else:
                raise ValidationError(f"Неподдерживаемый формат экспорта: {format}")

            return file_content_bytes, filename

        except NotFoundException:
            raise  # Передаем NotFoundException выше
        except ValidationError:
            raise  # Передаем ValidationError выше
        except Exception as e:
            logger.error(f"Общая ошибка при экспорте курса {course_id}: {str(e)}")
            # Не откатываем транзакцию, т.к. мы только читали данные
            raise  # Передаем другие исключения

    # Методы работы с шаблонами уроков
    @memory_optimized()
    async def get_lesson_templates(self, type: Optional[str] = None) -> List[LessonTemplate]:
        """Получает шаблоны уроков с оптимизацией запросов"""
        try:
            # Проверяем кэш
            cache_key = f"lesson_templates:{type or 'all'}"
            cached_templates = await self.cache_service.get_cached_data(cache_key)
            if cached_templates:
                return cached_templates

            # Формируем оптимизированный запрос
            query = select(LessonTemplate)
            if type:
                query = query.where(LessonTemplate.type == type)

            query = await self.query_optimizer.optimize_query(query)
            result = await self.session.execute(query)
            templates = result.scalars().all()

            # Кэшируем результат
            await self.cache_service.cache_data(cache_key, templates, ttl=3600)

            return templates
        except Exception as e:
            logger.error(f"Error getting lesson templates: {str(e)}")
            raise

    async def create_lesson_template(self, template_data: Dict[str, Any]) -> LessonTemplate:
        """Создаёт новый шаблон урока"""
        try:
            template = LessonTemplate(
                name=template_data['name'],
                type=template_data['type'],
                structure=template_data['structure'],
                is_default=template_data.get('is_default', False)
            )
            self.session.add(template)
            await self.session.commit()
            await self.session.refresh(template)

            # Инвалидируем кэш шаблонов
            await self.cache_service.invalidate_pattern("lesson_templates:*")

            return template
        except Exception as e:
            logger.error(f"Error creating lesson template: {str(e)}")
            await self.session.rollback()
            raise

    @memory_optimized()
    async def update_lesson_template(
            self,
            template_id: int,
            template_data: Dict[str, Any]
    ) -> Optional[LessonTemplate]:
        """Обновляет существующий шаблон урока"""
        try:
            query = await self.query_optimizer.optimize_query(
                select(LessonTemplate).where(LessonTemplate.id == template_id)
            )
            result = await self.session.execute(query)
            template = result.scalar_one_or_none()

            if not template:
                return None

            for field, value in template_data.items():
                setattr(template, field, value)

            await self.session.commit()
            await self.session.refresh(template)

            # Инвалидируем кэш шаблонов
            await self.cache_service.invalidate_pattern("lesson_templates:*")

            return template
        except Exception as e:
            logger.error(f"Error updating lesson template: {str(e)}")
            await self.session.rollback()
            raise

    async def delete_lesson_template(self, template_id: int) -> bool:
        """Удаляет шаблон урока"""
        try:
            query = await self.query_optimizer.optimize_query(
                select(LessonTemplate).where(LessonTemplate.id == template_id)
            )
            result = await self.session.execute(query)
            template = result.scalar_one_or_none()

            if not template:
                return False

            await self.session.delete(template)
            await self.session.commit()

            # Инвалидируем кэш шаблонов
            await self.cache_service.invalidate_pattern("lesson_templates:*")

            return True
        except Exception as e:
            logger.error(f"Error deleting lesson template: {str(e)}")
            await self.session.rollback()
            raise

    @memory_optimized()
    async def get_course_templates(
            self,
            level: Optional[str] = None,
            target_audience: Optional[str] = None
    ) -> List[Dict[str, Any]]:
        """Получает шаблоны курсов с фильтрацией и оптимизацией"""
        try:
            # Проверяем кэш
            cache_key = f"course_templates:{level or 'all'}:{target_audience or 'all'}"
            cached_templates = await self.cache_service.get_cached_data(cache_key)
            if cached_templates:
                return cached_templates

            # Формируем оптимизированный запрос
            query = select(Course).where(Course.is_template == True)

            if level:
                query = query.where(Course.level == level)
            if target_audience:
                query = query.where(Course.target_audience == target_audience)

            query = await self.query_optimizer.optimize_query(query)
            result = await self.session.execute(query)
            templates = result.scalars().all()

            # Преобразуем в словари
            templates_data = [
                {
                    "id": template.id,
                    "name": template.name,
                    "level": template.level,
                    "target_audience": template.target_audience,
                    "format": template.format,
                    "total_duration": template.total_duration,
                    "description": template.description
                }
                for template in templates
            ]

            # Кэшируем результат
            await self.cache_service.cache_data(cache_key, templates_data, ttl=3600)

            return templates_data
        except Exception as e:
            logger.error(f"Error getting course templates: {str(e)}")
            raise

    async def mark_as_template(self, course_id: int) -> bool:
        """Помечает курс как шаблон"""
        try:
            query = await self.query_optimizer.optimize_query(
                select(Course).where(Course.id == course_id)
            )
            result = await self.session.execute(query)
            course = result.scalar_one_or_none()

            if not course:
                return False

            course.is_template = True
            await self.session.commit()

            # Инвалидируем кэш шаблонов
            await self.cache_service.invalidate_pattern("course_templates:*")

            return True
        except Exception as e:
            logger.error(f"Error marking course as template: {str(e)}")
            await self.session.rollback()
            raise

    @memory_optimized()
    async def generate_course_from_template(
            self,
            template_id: int,
            new_name: str,
            user_id: int,
            customization: Dict[str, Any] = None
    ) -> Course:
        """Генерирует новый курс на основе шаблона с оптимизацией"""
        try:
            template = await self.get_course(template_id)
            if not template or not template.is_template:
                raise NotFoundException("Template not found")

            # Создаём новый курс на основе шаблона с учётом кастомизации
            course_data = {
                "name": new_name,
                "language": template.language,
                "level": template.level,
                "target_audience": template.target_audience,
                "format": template.format,
                "description": template.description,
                "exam_prep": template.exam_prep,
                "prerequisites": template.prerequisites.copy(),
                "learning_outcomes": template.learning_outcomes.copy(),
                "creator_id": user_id,
                "is_template": False,
                "is_used": False,
                "usage_count": 0
            }

            # Применяем кастомизацию
            if customization:
                course_data.update(customization)

            # Set initial total_duration from template
            initial_total_duration = template.total_duration

            # Создаём курс
            new_course = Course(
                **course_data,
                total_duration=initial_total_duration  # Set initial total_duration
            )
            self.session.add(new_course)
            await self.session.flush()

            # Копируем структуру из шаблона
            await self._copy_course_structure(template, new_course)

            # Recalculate total_duration after copying structure
            query = await self.query_optimizer.optimize_query(
                select(Lesson).where(Lesson.course_id == new_course.id)
            )
            result = await self.session.execute(query)
            lessons = result.scalars().all()

            # Update total_duration based on actual lesson durations
            recalculated_total_duration = sum(lesson.duration for lesson in lessons)
            new_course.total_duration = recalculated_total_duration

            await self.session.commit()
            await self.session.refresh(new_course)

            return new_course

        except Exception as e:
            logger.error(f"Error generating course from template: {str(e)}")
            await self.session.rollback()
            raise

    async def _copy_course_structure(self, source: Course, target: Course):
        """Копирует структуру курса из шаблона с батч-процессингом"""
        try:
            # Получаем все уроки шаблона оптимизированным запросом
            query = await self.query_optimizer.optimize_query(
                select(Lesson).where(Lesson.course_id == source.id)
            )
            result = await self.session.execute(query)
            source_lessons = result.scalars().all()

            # Подготавливаем данные для батч-вставки
            lessons_data = []
            activities_data = []

            for lesson in source_lessons:
                new_lesson = Lesson(
                    course_id=target.id,
                    title=lesson.title,
                    duration=lesson.duration,
                    order=lesson.order,
                    objectives=lesson.objectives.copy(),
                    grammar=lesson.grammar.copy(),
                    vocabulary=lesson.vocabulary.copy(),
                    materials=lesson.materials.copy(),
                    homework=lesson.homework.copy()
                )
                self.session.add(new_lesson)
                await self.session.flush()
                lessons_data.append(new_lesson)

                # Копируем активности
                activity_query = await self.query_optimizer.optimize_query(
                    select(Activity).where(Activity.lesson_id == lesson.id)
                )
                activity_result = await self.session.execute(activity_query)
                activities = activity_result.scalars().all()

                for activity in activities:
                    new_activity = Activity(
                        lesson_id=new_lesson.id,
                        name=activity.name,
                        type=activity.type,
                        duration=activity.duration,
                        description=activity.description,
                        materials=activity.materials.copy(),
                        objectives=activity.objectives.copy()
                    )
                    activities_data.append(new_activity)

            # Батч-вставка всех активностей
            await self.batch_processor.bulk_insert(activities_data)

        except Exception as e:
            logger.error(f"Error copying course structure: {str(e)}")
            raise

    @memory_optimized()
    async def generate_next_batch(
        self,
        course_id: int,
        current_lesson_count: int,
        user_id: int, # Добавим user_id для возможного отслеживания
        batch_size: int = 1 # Размер батча по умолчанию, можно сделать настраиваемым
    ) -> List[Lesson]:
        """Генерирует и сохраняет следующую порцию уроков для курса."""
        logger.info(f"Запрос на генерацию следующей ({batch_size}) порции уроков для курса {course_id} после {current_lesson_count} уроков.")
        import json # Добавим импорт json здесь, если он нужен локально
        import traceback # Добавим импорт traceback здесь

        try:
            # 1. Получаем курс и проверяем существование
            query = await self.query_optimizer.optimize_query(
                select(Course).where(Course.id == course_id)
            )
            result = await self.session.execute(query)
            course = result.scalar_one_or_none()

            if not course:
                raise NotFoundException(f"Курс с ID {course_id} не найден.")

            # 2. Определяем диапазон для нового батча
            start_lesson_num = current_lesson_count + 1
            total_lessons_planned = getattr(course, 'lessons_count', current_lesson_count + batch_size)
            end_lesson_num = min(start_lesson_num + batch_size - 1, total_lessons_planned)

            if start_lesson_num > end_lesson_num:
                logger.warning(f"Для курса {course_id} уже сгенерированы все ({total_lessons_planned}) запланированные уроки.")
                return []

            lessons_to_generate_count = end_lesson_num - start_lesson_num + 1
            logger.info(f"Планируется генерация уроков {start_lesson_num}-{end_lesson_num} ({lessons_to_generate_count} шт.)")

            # 3. Получаем контекст предыдущих уроков (последние N)
            previous_lessons_context = []
            if current_lesson_count > 0:
                context_limit = 3
                # Импортируем selectinload здесь, если он нужен
                from sqlalchemy.orm import selectinload
                context_query = await self.query_optimizer.optimize_query(
                    select(Lesson)
                    .where(Lesson.course_id == course_id)
                    .order_by(Lesson.order.desc())
                    .limit(context_limit)
                    # .options(selectinload(Lesson.activities)) # Можно раскомментировать для большего контекста
                )
                context_result = await self.session.execute(context_query)
                previous_lessons_context = [
                    {
                        "title": l.title,
                        "objectives": l.objectives,
                        "grammar": l.grammar,
                        "vocabulary": l.vocabulary
                    }
                    for l in reversed(context_result.scalars().all())
                ]
                logger.info(f"Получен контекст из {len(previous_lessons_context)} предыдущих уроков.")

            # 4. Подготавливаем данные/промпт для AI
            # Используем данные из существующего объекта course
            # Убедимся, что все необходимые поля присутствуют в CourseCreate или извлекаются из course
            batch_course_data_dict = {
                "name": course.name,
                "language": course.language,
                "level": course.level,
                "start_level": getattr(course, 'start_level', course.level),
                "target_audience": course.target_audience,
                "format": course.format,
                "description": course.description,
                "exam_prep": course.exam_prep,
                "prerequisites": course.prerequisites,
                "learning_outcomes": course.learning_outcomes,
                "lessons_count": lessons_to_generate_count,
                "lesson_duration": getattr(course, 'lesson_duration', 60),
                "methodology": getattr(course, 'methodology', 'communicative'),
                "main_topics": getattr(course, 'main_topics', ""),
                "grammar_focus": getattr(course, 'grammar_focus', ""),
                "vocabulary_focus": getattr(course, 'vocabulary_focus', ""),
                "student_age": getattr(course, 'student_age', ""),
                "student_interests": getattr(course, 'student_interests', ""),
                "student_goals": getattr(course, 'student_goals', ""),
                "common_mistakes": getattr(course, 'common_mistakes', ""),
                "include_speaking": getattr(course, 'include_speaking', True),
                "include_listening": getattr(course, 'include_listening', True),
                "include_reading": getattr(course, 'include_reading', True),
                "include_writing": getattr(course, 'include_writing', True),
                "include_games": getattr(course, 'include_games', True),
                "custom_exam": getattr(course, 'custom_exam', ""),
                "exam_prep_lessons": getattr(course, 'exam_prep_lessons', 0),
                "lessons": [] # Важно: не передаем существующие уроки
            }
            # Создаем объект CourseCreate, пропуская None значения, если модель этого требует
            batch_course_data = CourseCreate(**{k: v for k, v in batch_course_data_dict.items() if v is not None})


            prompt = self._create_ai_prompt(batch_course_data)

            prompt['batch_info'] = {
                'current_batch_start': start_lesson_num,
                'current_batch_end': end_lesson_num,
                'total_lessons_planned': total_lessons_planned
            }
            if previous_lessons_context:
                prompt['previous_lessons_context'] = previous_lessons_context

            prompt['requirements'] = f"""
            {prompt.get('requirements', '')}

            Важно: Это запрос на генерацию уроков с {start_lesson_num} по {end_lesson_num} для курса, который уже содержит {current_lesson_count} уроков.
            Предыдущие уроки (контекст): {json.dumps(previous_lessons_context, ensure_ascii=False)[:500]}...
            Пожалуйста, сгенерируй ТОЛЬКО уроки с {start_lesson_num} по {end_lesson_num}.
            Убедись, что новые уроки логически продолжают предыдущие, учитывая пройденные темы, грамматику и лексику из контекста.
            Сохраняй общую структуру и методику курса.
            """

            # 5. Вызываем AI сервис
            logger.info(f"Отправка запроса в AI для генерации уроков {start_lesson_num}-{end_lesson_num}")
            # Убедимся, что метод ai_service существует и принимает такой промпт
            generated_structure = await self.ai_service.generate_course_structure(prompt)

            # 6. Валидация и обработка ответа AI
            self._validate_generated_structure(generated_structure)
            new_lessons_data = generated_structure.get('lessons', [])

            if not new_lessons_data:
                logger.warning("AI не вернул уроки для нового батча.")
                return []

            logger.info(f"AI вернул {len(new_lessons_data)} уроков.")

            # 7. Сохранение новых уроков и активностей в БД
            created_lessons: List[Lesson] = []
            total_new_duration = 0

            for lesson_index, lesson_structure in enumerate(new_lessons_data):
                lesson_order = start_lesson_num + lesson_index
                lesson = Lesson(
                    course_id=course.id,
                    title=lesson_structure.get('title', f'Урок {lesson_order}'),
                    order=lesson_order,
                    objectives=lesson_structure.get('objectives', []),
                    grammar=lesson_structure.get('grammar', []),
                    vocabulary=lesson_structure.get('vocabulary', []),
                    materials=lesson_structure.get('materials', []),
                    homework=lesson_structure.get('homework', {}),
                    duration=lesson_structure.get('duration', 60),
                    is_completed=False
                )
                self.session.add(lesson)
                await self.session.flush()

                lesson_duration = 0
                activities_to_insert = []
                for activity_index, activity_data in enumerate(lesson_structure.get('activities', [])):
                    activity = Activity(
                        lesson_id=lesson.id,
                        name=activity_data.get('name', f'Активность {activity_index+1}'),
                        type=activity_data.get('type', 'practice'),
                        duration=activity_data.get('duration', 15),
                        description=activity_data.get('description', ''),
                        materials=activity_data.get('materials', []),
                        objectives=activity_data.get('objectives', [])
                    )
                    activities_to_insert.append(activity)
                    lesson_duration += activity.duration

                await self.batch_processor.bulk_insert(activities_to_insert)
                lesson.duration = lesson_duration
                created_lessons.append(lesson)
                total_new_duration += lesson_duration

            course.total_duration = (course.total_duration or 0) + total_new_duration
            await self.session.commit()

            logger.info(f"Успешно создано и сохранено {len(created_lessons)} новых уроков.")

            # Возвращаем только что созданные уроки
            # Можно сделать refresh или отдельный запрос для получения полных данных
            return created_lessons

        except NotFoundException:
            logger.warning(f"Курс {course_id} не найден при попытке генерации следующей части.")
            raise
        except ValidationError as ve:
             logger.error(f"Ошибка валидации при генерации следующей части: {ve}")
             await self.session.rollback()
             # Импортируем HTTPException и status, если они еще не импортированы вверху файла
             from fastapi import HTTPException, status
             raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=str(ve))
        except Exception as e:
            logger.error(f"Непредвиденная ошибка при генерации следующей части уроков для курса {course_id}: {str(e)}")
            logger.error(traceback.format_exc())
            await self.session.rollback()
            raise

    async def __aenter__(self):
        return self

    async def __aexit__(self, exc_type, exc_val, exc_tb):
        await self.session.close()
